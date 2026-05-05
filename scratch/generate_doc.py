
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_doc():
    doc = docx.Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    with open('FINAL_PROJECT_WRITEUP.md', 'r') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
            
        if line.startswith('# '):
            p = doc.add_heading(line[2:], 0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], 1)
        elif line.startswith('|'):
            # Table detected
            # Read header
            header = [cell.strip() for cell in line.split('|') if cell.strip()]
            i += 1 # Skip separator line |---|---|
            i += 1
            
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row = [cell.strip().replace('**', '') for cell in lines[i].split('|') if cell.strip()]
                rows.append(row)
                i += 1
            
            table = doc.add_table(rows=1, cols=len(header))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for idx, text in enumerate(header):
                hdr_cells[idx].text = text
                
            for row_data in rows:
                row_cells = table.add_row().cells
                for idx, text in enumerate(row_data):
                    row_cells[idx].text = text
            doc.add_paragraph()
            continue # Already advanced i
        elif line.startswith('*'):
            p = doc.add_paragraph(line[1:].strip(), style='List Bullet')
        elif line.startswith('!['):
            # Image
            # Extract path: ![alt](path)
            start = line.find('(') + 1
            end = line.find(')')
            img_path = line[start:end]
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(6))
            else:
                doc.add_paragraph(f"[Image Missing: {img_path}]")
        elif line == '---':
            pass # Skip horizontal rules
        else:
            doc.add_paragraph(line)
            
        i += 1

    doc.save('FINAL_PROJECT_WRITEUP.docx')
    print("Document saved as FINAL_PROJECT_WRITEUP.docx")

if __name__ == '__main__':
    create_doc()
